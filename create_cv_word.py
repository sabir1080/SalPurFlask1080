from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# HEADER
header = doc.add_paragraph()
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
header_run = header.add_run("SABIR SHAH")
header_run.font.size = Pt(28)
header_run.font.bold = True
header_run.font.color.rgb = RGBColor(30, 60, 114)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sub_run = subtitle.add_run("Business Software Developer")
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(42, 82, 152)

sub_run2 = subtitle.add_run("\nPython | Flask | PostgreSQL | Double-Entry Accounting Systems")
sub_run2.font.size = Pt(11)
sub_run2.font.color.rgb = RGBColor(52, 73, 94)

# Contact info
contact = doc.add_paragraph()
contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact_run = contact.add_run("Karachi, Pakistan | +92 345 3231545 | sabir1212temp@yahoo.com | github.com/sabir1080")
contact_run.font.size = Pt(10)

doc.add_paragraph()  # spacing

# PROFESSIONAL SUMMARY
heading = doc.add_heading('PROFESSIONAL SUMMARY', level=1)
heading.runs[0].font.size = Pt(12)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

summary_text = "Results-driven Business Software Developer with 20+ years of proven expertise designing, building, and deploying enterprise-grade inventory management, accounting systems, and point-of-sale solutions. Demonstrates deep mastery of double-entry accounting principles, weighted-average inventory costing methodologies, and comprehensive financial reporting. Specializes in architecting scalable web applications using Python, Flask, and PostgreSQL with a proven track record of delivering clean, well-tested code backed by 120+ automated tests and continuous integration pipelines. Recognized for honest project scoping, meticulous attention to data integrity, and consistent delivery of production-ready solutions."

p = doc.add_paragraph(summary_text)
p.paragraph_format.space_before = Pt(0)
for run in p.runs:
    run.font.size = Pt(10)

# CORE COMPETENCIES
heading = doc.add_heading('CORE COMPETENCIES', level=1)
heading.runs[0].font.size = Pt(12)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

skills = [
    ("Languages & Frameworks", "Python 3, Flask, SQLAlchemy ORM, JavaScript ES6+, HTML5, CSS3, Bootstrap 5, Jinja2 templating"),
    ("Databases & Cloud Infrastructure", "PostgreSQL, MySQL, SQLite, Docker containerization, CI/CD (GitHub Actions), Gunicorn, Render, Neon cloud databases"),
    ("Engineering Best Practices", "Automated testing (pytest), RESTful API design, Application security (Argon2 password hashing, CSRF protection), Git & GitHub workflows"),
    ("Domain Expertise", "Double-entry accounting, Inventory costing methods, Invoice management, Payroll processing, Fixed asset accounting, Financial statement generation, POS systems, Audit trails"),
]

for skill_title, skill_content in skills:
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(skill_title + ": ")
    p_run.bold = True
    p_run.font.size = Pt(10)
    p.add_run(skill_content).font.size = Pt(10)

# PROFESSIONAL EXPERIENCE
heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
heading.runs[0].font.size = Pt(12)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

# Job 1
p = doc.add_paragraph()
run = p.add_run("Business Software Developer")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(30, 60, 114)
p.add_run(" | Independent · Karachi | Present").font.size = Pt(10)

job1_points = [
    "Architected and built a production-ready trading and accounting system featuring integrated point-of-sale, general ledger, inventory management, and comprehensive financial reporting.",
    "Implemented comprehensive double-entry general ledger with control accounts, accounting periods, year-end closing procedures, and weighted-average inventory costing.",
    "Developed advanced point-of-sale module with barcode/QR code scanning, multi-payment method processing, customer management, and thermal receipt printing.",
    "Built complete suite of financial reports including Balance Sheet, Profit & Loss statements, Trial Balances, Cash Flow analysis, and automated Reconciliation reports.",
    "Engineered immutable document management system with full audit trail including timestamps and user attribution, role-based access control, and one-click backup/restore functionality.",
    "Deployed to cloud infrastructure (PostgreSQL + Render) with continuous integration/continuous deployment pipeline and zero-downtime deployment strategy.",
]

for point in job1_points:
    p = doc.add_paragraph(point, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# Job 2
p = doc.add_paragraph()
run = p.add_run("Senior Software Developer")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(30, 60, 114)
p.add_run(" | Independent · Karachi | Previous").font.size = Pt(10)

job2_points = [
    "Developed specialized systems for inventory management, sales & purchase processing, payroll, HR, transportation, garment production, and oil tanker operations.",
    "Maintained systems in daily production use with on-site support, continuous optimization, and proactive troubleshooting.",
    "Built applications with relational databases (SQL) emphasizing data integrity, transactional consistency, and regulatory compliance.",
    "Established strong long-term client relationships through clear technical communication, domain expertise, and commitment to solving real business problems.",
]

for point in job2_points:
    p = doc.add_paragraph(point, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# FLAGSHIP PROJECT
heading = doc.add_heading('FLAGSHIP PROJECT: TRADEFLOW', level=1)
heading.runs[0].font.size = Pt(12)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

p = doc.add_paragraph("Complete Inventory & Double-Entry Accounting System")
p_run = p.runs[0]
p_run.bold = True
p_run.font.size = Pt(11)

p = doc.add_paragraph("A sophisticated, production-deployed web application that unifies inventory management, sales/purchase operations, and financial accounting through a single double-entry general ledger.")
for run in p.runs:
    run.font.size = Pt(10)

project_points = [
    "Accounting Core: Full double-entry general ledger with control accounts, accounting periods, automatic year-end closing, and gapless invoice numbering.",
    "Inventory Management: Weighted-average costing, real-time stock tracking, barcode/QR labels, multi-location warehouse support, and automated reorder alerts.",
    "Point-of-Sale: Barcode scanning, multiple payment methods, customer management, quick item lookup, and thermal receipt printing.",
    "Fixed Assets: Complete fixed asset register with straight-line and reducing-balance depreciation methods and lifecycle tracking.",
    "Financial Reporting: Balance sheets, profit & loss statements, trial balances, cash flow analysis, and reconciliation reports.",
    "Security & Auditability: Role-based access control, immutable transaction records, comprehensive audit trail with timestamps.",
    "Configurability: Timezone support, fiscal-year customization, multi-currency capability for global deployment.",
    "Technical Stack: Python 3 / Flask / SQLAlchemy ORM / PostgreSQL / Bootstrap 5 / Jinja2. 120+ pytest-based automated tests.",
]

for point in project_points:
    p = doc.add_paragraph(point, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# EDUCATION
heading = doc.add_heading('EDUCATION & PROFESSIONAL CERTIFICATION', level=1)
heading.runs[0].font.size = Pt(12)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

education = [
    ("Certified Python Programmer", "Aptech, Karachi", "2025–2026"),
    ("Diploma in Computer Science", "Petroman Training Institute of Computer Sciences, Karachi", "1993–1994"),
    ("B.Sc. (Science)", "University of Karachi", "1991"),
]

for degree, institution, year in education:
    p = doc.add_paragraph()
    run = p.add_run(degree)
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(" | " + institution + " | " + year).font.size = Pt(10)

# PORTFOLIO
heading = doc.add_heading('PORTFOLIO & ONLINE PRESENCE', level=1)
heading.runs[0].font.size = Pt(12)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

portfolio = [
    ("GitHub", "github.com/sabir1080", "Source code and open-source contributions"),
    ("Live Demo", "TradeFlow-Demo.onrender.com", "Demo available: sabir.shah@tradeflowhq.com / demo1234"),
    ("LinkedIn", "linkedin.com/in/sabir-shah-19470b351", "Professional network and endorsements"),
    ("Fiverr", "fiverr.com/sabirshah7", "Freelance portfolio with client reviews"),
    ("YouTube", "@TradeFlowBusinessSolutions", "Technical tutorials and product walkthroughs"),
    ("Facebook", "Business Page & Community Forum", "Business updates and community engagement"),
    ("Direct Contact", "+92 345 3231545 (WhatsApp)", "Quick consultation and project inquiries"),
]

for label, link, desc in portfolio:
    p = doc.add_paragraph()
    run = p.add_run(label + ": ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(link).font.size = Pt(10)
    p.add_run(" | " + desc).font.size = Pt(9)
    p.runs[-1].italic = True

doc.save('SABIR_SHAH_CV.docx')
print("[OK] Professional CV Word file created successfully!")
print("[OK] File: SABIR_SHAH_CV.docx")
