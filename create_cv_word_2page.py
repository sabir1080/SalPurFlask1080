from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_horizontal_line(paragraph):
    """Add a horizontal line under paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2A5298')
    pBdr.append(bottom)
    pPr.append(pBdr)

doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# ==================== PAGE 1 ====================

# HEADER
header = doc.add_paragraph()
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
header_run = header.add_run("SABIR SHAH")
header_run.font.size = Pt(20)
header_run.font.bold = True
header_run.font.color.rgb = RGBColor(30, 60, 114)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sub_run = subtitle.add_run("Business Software Developer")
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(42, 82, 152)

sub_run2 = subtitle.add_run("\nPython | Flask | PostgreSQL | Double-Entry Accounting Systems")
sub_run2.font.size = Pt(9)
sub_run2.font.color.rgb = RGBColor(52, 73, 94)

# Contact info
contact = doc.add_paragraph()
contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact_run = contact.add_run("Karachi, Pakistan | +92 345 3231545 | sabir1212temp@yahoo.com | github.com/sabir1080")
contact_run.font.size = Pt(8.5)
add_horizontal_line(contact)

# PROFESSIONAL SUMMARY
heading = doc.add_heading('PROFESSIONAL SUMMARY', level=1)
heading.runs[0].font.size = Pt(10)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

summary_text = "Results-driven Business Software Developer with 20+ years of proven expertise designing, building, and deploying enterprise-grade inventory management, accounting systems, and point-of-sale solutions. Demonstrates deep mastery of double-entry accounting principles, weighted-average inventory costing methodologies, and comprehensive financial reporting. Specializes in architecting scalable web applications using Python, Flask, and PostgreSQL with a proven track record of delivering clean, well-tested code backed by 120+ automated tests and continuous integration pipelines. Recognized for honest project scoping, meticulous attention to data integrity, and consistent delivery of production-ready solutions."

p = doc.add_paragraph(summary_text)
for run in p.runs:
    run.font.size = Pt(9)

# CORE COMPETENCIES
heading = doc.add_heading('CORE COMPETENCIES', level=1)
heading.runs[0].font.size = Pt(10)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

skills = [
    ("Languages & Frameworks", "Python 3, Flask, SQLAlchemy ORM, JavaScript, HTML5, CSS3, Bootstrap 5, Jinja2"),
    ("Databases & Cloud Infrastructure", "PostgreSQL, MySQL, SQLite, Docker, CI/CD (GitHub Actions), Gunicorn, Render, Neon"),
    ("Engineering Best Practices", "Automated testing (pytest), RESTful API design, Security (Argon2 hashing, CSRF), Git & GitHub"),
    ("Domain Expertise", "Double-entry accounting, Inventory costing, Invoicing, Payroll, Fixed assets, Financial reporting, POS systems"),
]

for skill_title, skill_content in skills:
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(skill_title + ": ")
    p_run.bold = True
    p_run.font.size = Pt(9)
    p.add_run(skill_content).font.size = Pt(8.5)

# PROFESSIONAL EXPERIENCE
heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
heading.runs[0].font.size = Pt(10)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

# Job 1
p = doc.add_paragraph()
run = p.add_run("Business Software Developer")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(30, 60, 114)
p.add_run(" | Independent · Karachi | Present").font.size = Pt(9)

job1_points = [
    "Architected and built a production-ready trading and accounting system featuring integrated point-of-sale, general ledger, inventory management, and comprehensive financial reporting—deployed live with 120+ automated test coverage.",
    "Implemented comprehensive double-entry general ledger with control accounts, accounting periods, year-end closing, and weighted-average inventory costing.",
    "Developed advanced point-of-sale module with barcode/QR code scanning, multi-payment processing, customer management, and thermal receipt printing.",
    "Built complete suite of financial reports including Balance Sheet, P&L, Trial Balance, Cash Flow, and Reconciliation reports.",
    "Engineered immutable document management system with full audit trail, role-based access control, and one-click backup/restore.",
    "Deployed to cloud infrastructure (PostgreSQL + Render) with CI/CD pipeline and zero-downtime deployment strategy.",
]

for point in job1_points:
    p = doc.add_paragraph(point, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(8.5)

# Job 2
p = doc.add_paragraph()
run = p.add_run("Senior Software Developer")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(30, 60, 114)
p.add_run(" | Independent · Karachi | Previous").font.size = Pt(9)

job2_points = [
    "Developed specialized systems for inventory management, sales & purchase processing, payroll, HR, transportation, garment production, and oil tanker operations.",
    "Maintained systems in daily production use with on-site support, continuous optimization, and proactive troubleshooting.",
    "Built applications with relational databases (SQL) emphasizing data integrity, transactional consistency, and compliance.",
    "Established strong long-term client relationships through clear communication, domain expertise, and real business problem-solving.",
]

for point in job2_points:
    p = doc.add_paragraph(point, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(8.5)

# FLAGSHIP PROJECT - PAGE 1
heading = doc.add_heading('FLAGSHIP PROJECT: TRADEFLOW', level=1)
heading.runs[0].font.size = Pt(10)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

p = doc.add_paragraph("Complete Inventory & Double-Entry Accounting System")
p_run = p.runs[0]
p_run.bold = True
p_run.font.size = Pt(10)

p = doc.add_paragraph("A sophisticated, production-deployed web application that unifies inventory management, sales/purchase operations, and financial accounting through a single double-entry general ledger.")
for run in p.runs:
    run.font.size = Pt(9)

project_points_p1 = [
    "Accounting Core: Full double-entry general ledger with control accounts, accounting periods, automatic year-end closing.",
    "Inventory Management: Weighted-average costing, real-time stock tracking, barcode/QR labels, multi-location support.",
    "Point-of-Sale: Barcode scanning, multi-payment methods, customer management, quick item lookup, thermal receipt printing.",
    "Financial Reporting: Balance Sheet, P&L, Trial Balance, Cash Flow, and Reconciliation reports.",
]

for point in project_points_p1:
    p = doc.add_paragraph(point, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(8.5)

# PAGE BREAK
doc.add_page_break()

# ==================== PAGE 2 ====================

# FLAGSHIP PROJECT CONTINUED
heading = doc.add_heading('FLAGSHIP PROJECT: TRADEFLOW (CONTINUED)', level=1)
heading.runs[0].font.size = Pt(10)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

project_points_p2 = [
    "Fixed Assets: Complete asset register with straight-line and reducing-balance depreciation, lifecycle tracking.",
    "Security & Auditability: Role-based access control, immutable records, comprehensive audit trail with timestamps.",
    "Configurability: Timezone support, fiscal-year customization, multi-currency for global deployment.",
    "Technical Stack: Python 3 / Flask / SQLAlchemy ORM / PostgreSQL / Bootstrap 5 / Jinja2. Deployed to Render with GitHub Actions CI/CD. 120+ pytest automated tests.",
]

for point in project_points_p2:
    p = doc.add_paragraph(point, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(8.5)

# EDUCATION
heading = doc.add_heading('EDUCATION & PROFESSIONAL CERTIFICATION', level=1)
heading.runs[0].font.size = Pt(10)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

education = [
    ("Certified Python Programmer", "Aptech, Karachi", "2025–2026"),
    ("Diploma in Computer Science", "Petroman Training Institute of Computer Sciences, Karachi", "1993–1994"),
    ("B.Sc. (Science)", "University of Karachi", "1991"),
]

for degree, institution, year in education:
    p = doc.add_paragraph()
    run = p.add_run(degree)
    run.bold = True
    run.font.size = Pt(9)
    p.add_run(" | " + institution + " | " + year).font.size = Pt(8.5)

# PORTFOLIO
heading = doc.add_heading('PORTFOLIO & ONLINE PRESENCE', level=1)
heading.runs[0].font.size = Pt(10)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

portfolio = [
    ("GitHub", "github.com/sabir1080", "Source code and open-source contributions"),
    ("Live Demo", "TradeFlow-Demo.onrender.com", "Demo: sabir.shah@tradeflowhq.com / demo1234"),
    ("LinkedIn", "linkedin.com/in/sabir-shah-19470b351", "Professional network and endorsements"),
    ("Fiverr", "fiverr.com/sabirshah7", "Freelance portfolio with client reviews"),
    ("YouTube", "@TradeFlowBusinessSolutions", "Technical tutorials and product walkthroughs"),
    ("Facebook", "Business Page & Community Forum", "Business updates and community engagement"),
    ("Direct Contact", "+92 345 3231545 (WhatsApp)", "Quick consultation and project inquiries"),
]

for label, link, desc in portfolio:
    p = doc.add_paragraph()
    run = p.add_run(label + ": ")
    run.bold = True
    run.font.size = Pt(9)
    p.add_run(link).font.size = Pt(8.5)
    p.add_run(" | " + desc).font.size = Pt(8)
    p.runs[-1].italic = True

# TECHNICAL SKILLS SUMMARY
heading = doc.add_heading('ADDITIONAL TECHNICAL DETAILS', level=1)
heading.runs[0].font.size = Pt(10)
heading.runs[0].font.color.rgb = RGBColor(30, 60, 114)

details_text = "Testing & Quality: 120+ automated test cases with pytest, continuous integration via GitHub Actions. Database Design: Relational database design with referential integrity, transaction management, performance optimization. Security: Argon2 password hashing, CSRF protection, SQL injection prevention, role-based access control (RBAC). Deployment: Zero-downtime deployments, container orchestration, cloud database management, backup and disaster recovery procedures."

p = doc.add_paragraph(details_text)
for run in p.runs:
    run.font.size = Pt(8.5)

doc.save('CV_SABIR_SHAH_2page.docx')
print("[OK] 2-Page Professional CV created successfully!")
print("[OK] File: CV_SABIR_SHAH_2page.docx")
